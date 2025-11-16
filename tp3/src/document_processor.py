"""
Procesador de Documentos para Sistema RAG
========================================

Este módulo se encarga del procesamiento inteligente de documentos:
- Extracción de texto de múltiples formatos (PDF, TXT, DOCX)
- Chunking inteligente con RecursiveCharacterTextSplitter
- Metadata extraction y enriquecimiento
- Validación y limpieza de contenido

Autor: Sistema RAG Milei
Fecha: 2025-11-16
"""

import os
import re
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import hashlib

# Procesamiento de documentos
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

# Configuración de logging
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Procesador inteligente de documentos para el sistema RAG.
    
    Funcionalidades:
    - Extracción de texto de PDF, TXT, DOCX
    - Chunking inteligente con overlap controlado
    - Metadata extraction automática
    - Limpieza y normalización de texto
    - Validación de contenido
    """
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 min_chunk_size: int = 100,
                 max_chunk_size: int = 2000):
        """
        Inicializar procesador de documentos.
        
        Args:
            chunk_size: Tamaño óptimo de fragmentos
            chunk_overlap: Superposición entre fragmentos
            min_chunk_size: Tamaño mínimo de fragmento
            max_chunk_size: Tamaño máximo de fragmento
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        
        # Inicializar text splitter con configuración optimizada
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n",     # Párrafos
                "\n",       # Líneas
                ". ",       # Oraciones
                " ",        # Palabras
                ""          # Caracteres (último recurso)
            ],
            length_function=len,
            is_separator_regex=False
        )
        
        logger.info(f"📄 DocumentProcessor inicializado: chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def process_document(self, 
                        file_path: str, 
                        metadata: Optional[Dict] = None) -> List[LangChainDocument]:
        """
        Procesar un documento completo y dividirlo en chunks.
        
        Args:
            file_path: Ruta al archivo a procesar
            metadata: Metadata adicional del documento
            
        Returns:
            Lista de documentos fragmentados con metadata
        """
        try:
            # Verificar que el archivo existe
            if not os.path.exists(file_path):
                logger.error(f"Archivo no encontrado: {file_path}")
                return []
            
            # Extraer texto del archivo
            text_content = self._extract_text(file_path)
            if not text_content:
                logger.warning(f"No se pudo extraer texto de: {file_path}")
                return []
            
            # Limpiar y normalizar texto
            cleaned_text = self._clean_text(text_content)
            
            # Extraer metadata del documento
            doc_metadata = self._extract_metadata(file_path, metadata)
            
            # Crear documento LangChain
            document = LangChainDocument(
                page_content=cleaned_text,
                metadata=doc_metadata
            )
            
            # Dividir en chunks
            chunks = self.text_splitter.split_documents([document])
            
            # Enriquecer metadata de cada chunk
            enriched_chunks = self._enrich_chunk_metadata(chunks, doc_metadata)
            
            logger.info(f"✅ Documento procesado: {len(enriched_chunks)} chunks generados")
            
            return enriched_chunks
            
        except Exception as e:
            logger.error(f"Error procesando documento {file_path}: {e}")
            return []
    
    def _extract_text(self, file_path: str) -> Optional[str]:
        """
        Extraer texto de un archivo según su extensión.
        
        Args:
            file_path: Ruta del archivo
            
        Returns:
            Texto extraído o None si falla
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext == '.txt':
                return self._extract_txt_text(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            else:
                logger.warning(f"Formato de archivo no soportado: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"Error extrayendo texto de {file_path}: {e}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extraer texto de archivo PDF."""
        try:
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- PÁGINA {page_num + 1} ---\n"
                        text_content += page_text
            
            return text_content.strip() if text_content else None
            
        except Exception as e:
            logger.error(f"Error extrayendo PDF {file_path}: {e}")
            return None
    
    def _extract_txt_text(self, file_path: str) -> Optional[str]:
        """Extraer texto de archivo TXT."""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"No se pudo leer {file_path} con ninguna codificación")
            return None
            
        except Exception as e:
            logger.error(f"Error extrayendo TXT {file_path}: {e}")
            return None
    
    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extraer texto de archivo DOCX."""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error extrayendo DOCX {file_path}: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """
        Limpiar y normalizar texto extraído.
        
        Args:
            text: Texto a limpiar
            
        Returns:
            Texto limpio y normalizado
        """
        if not text:
            return ""
        
        # Normalizar saltos de línea
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Eliminar espacios múltiples
        text = re.sub(r' +', ' ', text)
        
        # Eliminar líneas vacías excesivas (más de 2 saltos)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Limpiar caracteres especiales problemáticos
        text = text.replace('\x00', '')  # Null bytes
        text = text.replace('\x0c', '')  # Form feeds
        
        # Normalizar puntuación
        text = re.sub(r'[ \t]+([.,;:!?])', r'\1', text)  # Espacio antes de puntuación
        
        # Eliminar líneas que son solo números o símbolos
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Mantener líneas que tienen contenido significativo
            if (len(stripped) > 3 and 
                not stripped.isdigit() and 
                not re.match(r'^[-=_*#]{3,}$', stripped)):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()
    
    def _extract_metadata(self, file_path: str, additional_metadata: Optional[Dict] = None) -> Dict[str, Any]:
        """
        Extraer metadata de un documento.
        
        Args:
            file_path: Ruta del archivo
            additional_metadata: Metadata adicional
            
        Returns:
            Dictionary con metadata del documento
        """
        file_path_obj = Path(file_path)
        
        # Metadata básica del archivo
        metadata = {
            'source': file_path,
            'filename': file_path_obj.name,
            'file_extension': file_path_obj.suffix.lower(),
            'file_size': os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            'processed_date': datetime.now().isoformat(),
            'doc_type': self._determine_doc_type(file_path_obj.name),
            'language': 'es',  # Asumimos español por defecto
            'encoding': 'utf-8'
        }
        
        # Extraer fecha del nombre del archivo si es posible
        date_from_filename = self._extract_date_from_filename(file_path_obj.name)
        if date_from_filename:
            metadata['extracted_date'] = date_from_filename
        
        # Agregar metadata adicional
        if additional_metadata:
            metadata.update(additional_metadata)
        
        # Generar hash único del documento
        try:
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
                metadata['content_hash'] = file_hash
        except Exception:
            pass
        
        return metadata
    
    def _determine_doc_type(self, filename: str) -> str:
        """Determinar el tipo de documento basado en el nombre del archivo."""
        filename_lower = filename.lower()
        
        if any(word in filename_lower for word in ['discurso', 'speech', 'alocución']):
            return 'discurso'
        elif any(word in filename_lower for word in ['entrevista', 'interview']):
            return 'entrevista'
        elif any(word in filename_lower for word in ['conferencia', 'press']):
            return 'conferencia_prensa'
        elif any(word in filename_lower for word in ['declaración', 'statement']):
            return 'declaracion'
        else:
            return 'documento'
    
    def _extract_date_from_filename(self, filename: str) -> Optional[str]:
        """Extraer fecha del nombre del archivo."""
        # Patrones de fecha comunes en nombres de archivo
        date_patterns = [
            r'(\d{4})-(\d{2})-(\d{2})',  # YYYY-MM-DD
            r'(\d{2})-(\d{2})-(\d{4})',  # DD-MM-YYYY
            r'(\d{4})(\d{2})(\d{2})',    # YYYYMMDD
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, filename)
            if match:
                try:
                    if len(match.group(1)) == 4:  # Año al inicio
                        year, month, day = match.groups()
                    else:  # DD-MM-YYYY
                        day, month, year = match.groups()
                    
                    # Validar fecha
                    date_str = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                    datetime.strptime(date_str, "%Y-%m-%d")
                    return date_str
                except ValueError:
                    continue
        
        return None
    
    def _enrich_chunk_metadata(self, 
                              chunks: List[LangChainDocument], 
                              base_metadata: Dict[str, Any]) -> List[LangChainDocument]:
        """
        Enriquecer metadata de cada chunk con información específica.
        
        Args:
            chunks: Lista de chunks de documentos
            base_metadata: Metadata base del documento
            
        Returns:
            Lista de chunks con metadata enriquecida
        """
        for i, chunk in enumerate(chunks):
            # Agregar metadata específica del chunk
            chunk.metadata.update({
                **base_metadata,
                'chunk_id': i + 1,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk.page_content),
                'chunk_start_char': 0 if i == 0 else sum(len(c.page_content) for c in chunks[:i])
            })
            
            # Agregar resumen del contenido
            chunk.metadata['content_preview'] = chunk.page_content[:100] + "..."
        
        return chunks
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Obtener estadísticas del procesador.
        
        Returns:
            Diccionario con estadísticas
        """
        return {
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "min_chunk_size": self.min_chunk_size,
            "max_chunk_size": self.max_chunk_size,
            "supported_formats": ['.pdf', '.txt', '.docx', '.doc'],
            "splitter_type": "RecursiveCharacterTextSplitter"
        }